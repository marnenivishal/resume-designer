#!/usr/bin/env python3
r"""Build a resume from a YAML source into PDF / HTML / DOCX / TXT.

One source of truth (resume.yaml) -> many targets. Rendering path:

    resume.yaml --(jinja2)--> HTML --(headless Chrome)--> PDF
                                   \--(python-docx)-----> DOCX   (ATS upload / recruiter edit)
                                   \--(plain text)------> TXT    (paste into web forms)

Usage:
    python build.py resume.yaml                       # -> PDF next to the source
    python build.py resume.yaml --format pdf,docx,txt
    python build.py resume.yaml --template classic --out ./build
    python build.py resume.yaml --check               # run ats_check.py on the PDF

Design rules live in DESIGN.md; the numbers come from assets/tokens.css.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

HERE = Path(__file__).resolve().parent
SKILL = HERE.parent
TEMPLATES = SKILL / "assets" / "templates"

# ---------------------------------------------------------------- dependencies

def _need(mod: str, pip: str | None = None):
    try:
        return __import__(mod)
    except ImportError:
        sys.exit(f"error: missing dependency '{mod}'. Install with: pip install {pip or mod}")


# ------------------------------------------------------------------- chrome

def find_chrome() -> str | None:
    """Locate a Chromium-family browser able to --print-to-pdf."""
    env = os.environ.get("RESUME_CHROME")
    if env and Path(env).exists():
        return env
    for name in ("chrome", "chromium", "chromium-browser", "google-chrome", "msedge"):
        p = shutil.which(name)
        if p:
            return p
    candidates = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
        "/usr/bin/google-chrome", "/usr/bin/chromium", "/usr/bin/chromium-browser",
        "/snap/bin/chromium",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    return None


def html_to_pdf(html_path: Path, pdf_path: Path, chrome: str) -> None:
    """Print HTML to PDF via headless Chrome.

    Two non-obvious requirements, both learned the hard way:
      * --headless=new is required; legacy --headless silently no-ops on some builds.
      * --print-to-pdf needs an ABSOLUTE path, or Chrome exits 0 writing nothing.
      * a unique --user-data-dir per invocation, or concurrent/serial runs block on
        the profile lock until they time out.
    """
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    if pdf_path.exists():
        pdf_path.unlink()
    udd = tempfile.mkdtemp(prefix="resume_chrome_")
    try:
        proc = subprocess.run(
            [chrome, "--headless=new", "--disable-gpu", "--no-sandbox",
             "--no-first-run", "--no-default-browser-check", "--disable-extensions",
             "--no-pdf-header-footer", "--run-all-compositor-stages-before-draw",
             "--virtual-time-budget=8000",
             f"--user-data-dir={udd}",
             f"--print-to-pdf={pdf_path.resolve()}",
             html_path.resolve().as_uri()],
            capture_output=True, text=True, timeout=180,
        )
    except subprocess.TimeoutExpired:
        sys.exit("error: Chrome timed out rendering the PDF")
    finally:
        shutil.rmtree(udd, ignore_errors=True)
    if not pdf_path.exists():
        sys.exit(f"error: Chrome produced no PDF.\n{proc.stderr[-2000:]}")


# -------------------------------------------------------------------- model

MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


def fmt_date(v, style: str = "short") -> str:
    """'2021-03' -> 'Mar 2021'; 'present' -> 'Present'; passthrough for free text."""
    if v in (None, ""):
        return ""
    s = str(v).strip()
    if s.lower() in ("present", "current", "now"):
        return "Present"
    m = re.fullmatch(r"(\d{4})-(\d{1,2})", s)
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12:
            return f"{MONTHS[mo - 1]} {y}" if style == "short" else f"{y}"
    if re.fullmatch(r"\d{4}", s):
        return s
    return s


def date_range(item: dict, style: str = "short") -> str:
    a, b = fmt_date(item.get("start"), style), fmt_date(item.get("end"), style)
    if a and b:
        return f"{a} – {b}"          # en dash
    return a or b or ""


def load_resume(path: Path) -> dict:
    yaml = _need("yaml", "pyyaml")
    with open(path, encoding="utf-8") as f:
        data = yaml.safe_load(f)
    if not isinstance(data, dict):
        sys.exit("error: resume file must be a YAML mapping")
    data.setdefault("config", {})
    data.setdefault("basics", {})
    if not data["basics"].get("name"):
        sys.exit("error: basics.name is required")
    return data


def validate(data: dict) -> list[str]:
    """Content lint. Returns human-readable warnings (never fatal)."""
    warn: list[str] = []
    b = data.get("basics", {})
    if not b.get("email"):
        warn.append("basics.email missing - recruiters need a reply path")
    loc = str(b.get("location", ""))
    if re.search(r"\d{1,5}\s+\w+\s+(St|Street|Ave|Avenue|Rd|Road|Blvd|Lane|Ln|Dr|Drive)\b", loc, re.I):
        warn.append("basics.location looks like a street address - use 'City, ST' only "
                    "(a full address adds no value and leaks PII)")

    WEAK = re.compile(r"\b(responsible for|duties includ\w*|worked on|helped with|"
                      r"assisted with|tasked with|in charge of)\b", re.I)
    FLUFF = re.compile(r"\b(team player|results[- ]driven|self[- ]starter|go[- ]getter|"
                       r"synergy|synergistic|think outside the box|hard[- ]working|"
                       r"detail[- ]oriented|proven track record|dynamic professional|"
                       r"passionate about)\b", re.I)
    for i, job in enumerate(data.get("experience") or []):
        for j, bullet in enumerate(job.get("bullets") or []):
            where = f"experience[{i}].bullets[{j}]"
            t = str(bullet)
            if WEAK.search(t):
                warn.append(f"{where}: weak opener ('{WEAK.search(t).group(0)}') - "
                            f"lead with an action verb and an outcome")
            if FLUFF.search(t):
                warn.append(f"{where}: cliche ('{FLUFF.search(t).group(0)}') - "
                            f"show it with evidence instead of claiming it")
            if not re.search(r"\d", t):
                warn.append(f"{where}: no number - quantify the outcome if you can "
                            f"(scale, %, time, money, volume)")
            if len(t) > 240:
                warn.append(f"{where}: {len(t)} chars - long bullets get skipped; aim under ~2 lines")
        if not job.get("bullets"):
            warn.append(f"experience[{i}] ({job.get('company','?')}): no bullets")
    if (data.get("summary") or "").strip() and len(str(data["summary"]).split()) > 70:
        warn.append("summary is long - 2-3 lines max, or cut it entirely")
    return warn


# ------------------------------------------------------------------- render

def build_env():
    j2 = _need("jinja2")
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATES)),
        autoescape=select_autoescape(["html", "xml"]),
        trim_blocks=True, lstrip_blocks=True,
    )
    env.filters["date"] = fmt_date
    env.filters["daterange"] = date_range
    return env


ALL_SECTIONS = ["experience", "projects", "education", "skills",
                "certifications", "publications", "awards"]

# Section order encodes ONE idea: the section carrying your strongest evidence goes
# first, because a reader who scans top-down may never reach the bottom. What counts
# as strongest evidence changes with career stage -- that is the only reason these
# differ.
STAGE_ORDER = {
    # No work history yet: education IS the credential, projects are the evidence.
    "student":   ["education", "projects", "experience", "skills", "certifications", "awards", "publications"],
    "new-grad":  ["education", "experience", "projects", "skills", "certifications", "awards", "publications"],
    # Once you have shipped things, work outranks the degree -- permanently.
    "mid":       ["experience", "projects", "skills", "education", "certifications", "awards", "publications"],
    "senior":    ["experience", "projects", "skills", "education", "certifications", "awards", "publications"],
    "executive": ["experience", "awards", "education", "skills", "certifications", "projects", "publications"],
    "academic":  ["education", "experience", "publications", "awards", "projects", "skills", "certifications"],
}


def infer_stage(data: dict) -> str:
    """Guess career stage from the data when config.stage is absent."""
    exp = data.get("experience") or []
    if data.get("publications"):
        return "academic"
    if not exp:
        return "student"
    years = 0
    for j in exp:
        s, e = str(j.get("start", "")), str(j.get("end", ""))
        ms, me = re.match(r"(\d{4})", s), re.match(r"(\d{4})", e)
        end_y = datetime.now().year if e.lower() in ("present", "current", "") else (int(me.group(1)) if me else 0)
        if ms and end_y:
            years += max(0, end_y - int(ms.group(1)))
    if years <= 1:
        return "new-grad"
    if years <= 7:
        return "mid"
    return "senior"


def section_order(data: dict) -> list[str]:
    cfg = data.get("config", {})
    if cfg.get("section_order"):
        order = [s for s in cfg["section_order"] if s in ALL_SECTIONS]
        return order + [s for s in ALL_SECTIONS if s not in order]
    stage = cfg.get("stage") or infer_stage(data)
    return STAGE_ORDER.get(stage, STAGE_ORDER["mid"])


def contact_items(data: dict) -> list[str]:
    """Contact line, in scan order. Label over raw URL: 'github.com/jane' beats
    'https://github.com/jane?tab=repositories'."""
    b = data.get("basics", {})
    out = [b.get("email"), b.get("phone"), b.get("location")]
    for l in (b.get("links") or []):
        if isinstance(l, dict):
            out.append(l.get("label") or re.sub(r"^https?://(www\.)?", "", str(l.get("url", ""))))
        else:
            out.append(re.sub(r"^https?://(www\.)?", "", str(l)))
    return [str(x).strip() for x in out if x]


def render_html(data: dict, template: str, overrides: dict | None = None) -> str:
    env = build_env()
    name = template if template.endswith(".html.j2") else f"{template}.html.j2"
    if not (TEMPLATES / name).exists():
        avail = ", ".join(sorted(p.stem.replace(".html", "") for p in TEMPLATES.glob("*.html.j2")
                                 if not p.name.startswith("_")))
        sys.exit(f"error: unknown template '{template}'. Available: {avail}")
    tokens = (SKILL / "assets" / "tokens.css").read_text(encoding="utf-8")
    cfg = data.get("config", {})
    ctx = dict(data)
    ctx.pop("config", None)
    html = env.get_template(name).render(
        **ctx,
        config=cfg,
        tokens_css=tokens,
        contact_items=contact_items(data),
        section_order=section_order(data),
        labels=cfg.get("labels", {}) or {},
        page=str(cfg.get("page", "letter")).lower(),
        accent=cfg.get("accent"),
        density=cfg.get("density", overrides.get("__density") if overrides else None) or cfg.get("density", "normal"),
        now=datetime.now(),
    )
    if overrides:
        css = "".join(f"{k}:{v};" for k, v in overrides.items() if not k.startswith("__"))
        if css:
            html = html.replace("</style>", f"\n/* fit pass */\n:root{{{css}}}\n</style>", 1)
        if overrides.get("__density"):
            html = re.sub(r'class="density-\w+"', f'class="density-{overrides["__density"]}"', html, count=1)
    return html


# The fit ladder. Each rung is a compromise a designer would actually accept; the
# ladder stops before the document starts looking cramped. Cramming a resume to fit
# is how it ends up unreadable -- past the last rung the honest answer is "cut a
# bullet", not "shrink it to 8pt".
FIT_LADDER = [
    ({}, "as configured"),
    ({"--lh-body": "1.28", "--gap-section": "12pt"}, "tightened leading and section gaps"),
    ({"__density": "tight"}, "density: tight"),
    ({"__density": "tight", "--page-margin-y": "0.5in", "--page-margin-x": "0.55in"},
     "density: tight + 0.5in margins"),
]


def pdf_pages(path: Path) -> int:
    try:
        import pypdf
        return len(pypdf.PdfReader(str(path)).pages)
    except Exception:
        return -1


def build_pdf_fitted(data: dict, template: str, html_path: Path, pdf_path: Path,
                     chrome: str, max_pages: int | None, quiet: bool) -> tuple[int, str | None]:
    """Render, and if it overflows max_pages walk the fit ladder.

    Returns (pages, note). Never silently shrinks: the note says what was done.
    """
    note = None
    for i, (overrides, label) in enumerate(FIT_LADDER):
        html_path.write_text(render_html(data, template, overrides), encoding="utf-8")
        html_to_pdf(html_path, pdf_path, chrome)
        n = pdf_pages(pdf_path)
        if not max_pages or n <= max_pages:
            if i > 0:
                note = f"fit to {n} page(s) via {label}"
                if not quiet:
                    print(f"  fit: overflowed {max_pages} page(s); applied {label}")
            return n, note
        if not quiet and i < len(FIT_LADDER) - 1:
            print(f"  fit: {n} pages at '{label}' — trying tighter…")
    n = pdf_pages(pdf_path)
    note = (f"still {n} pages at the tightest tasteful setting — this is a content "
            f"problem, not a layout one")
    if not quiet:
        print(f"  fit: {note}\n       cut bullets from older roles, or raise config.max_pages.")
    return n, note


def render_txt(data: dict) -> str:
    """Plain text for paste-into-form fields. No decoration, ASCII only."""
    out: list[str] = []
    b = data["basics"]
    out += [b["name"].upper(), ""]
    contact = [b.get("headline"), b.get("email"), b.get("phone"), b.get("location")]
    contact += [l.get("url") or l.get("label") for l in (b.get("links") or [])]
    out += [" | ".join(str(c) for c in contact if c), ""]
    if data.get("summary"):
        out += ["SUMMARY", str(data["summary"]).strip(), ""]
    if data.get("experience"):
        out.append("EXPERIENCE")
        for j in data["experience"]:
            out.append(f"{j.get('role','')} - {j.get('company','')}"
                       + (f" - {j['location']}" if j.get("location") else ""))
            out.append(date_range(j))
            out += [f"- {bl}" for bl in (j.get("bullets") or [])]
            out.append("")
    for key, title in (("projects", "PROJECTS"), ("education", "EDUCATION")):
        if data.get(key):
            out.append(title)
            for e in data[key]:
                head = e.get("name") or e.get("degree") or ""
                org = e.get("institution") or e.get("org") or ""
                out.append(" - ".join(x for x in (head, org, date_range(e)) if x))
                if e.get("summary"):
                    out.append(str(e["summary"]))
                out += [f"- {bl}" for bl in (e.get("bullets") or [])]
            out.append("")
    if data.get("skills"):
        out.append("SKILLS")
        for g in data["skills"]:
            if isinstance(g, dict):
                out.append(f"{g.get('group','')}: {', '.join(g.get('items') or [])}".strip(": "))
            else:
                out.append(str(g))
        out.append("")
    for sec in (data.get("custom_sections") or []):
        out.append(str(sec.get("title", "")).upper())
        for it in (sec.get("items") or []):
            out.append(f"- {it}" if not isinstance(it, dict) else f"- {it.get('text','')}")
        out.append("")
    text = "\n".join(out)
    text = text.replace("–", "-").replace("—", "-").replace("•", "-")
    text = text.replace("’", "'").replace("“", '"').replace("”", '"')
    return re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"


def render_docx(data: dict, out: Path) -> None:
    """Single-column DOCX built from real Word heading styles.

    Deliberately plainer than the PDF: this is the file you hand to a recruiter who
    wants to edit it, or upload where DOCX is requested. Structure over styling.
    """
    _need("docx", "python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cfg = data.get("config", {})
    accent = (cfg.get("accent") or "#1F4E5F").lstrip("#")
    rgb = RGBColor(int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16))

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(46)
        s.left_margin = s.right_margin = Pt(50)
    normal = doc.styles["Normal"]
    normal.font.name = cfg.get("docx_font", "Calibri")
    normal.font.size = Pt(10.5)

    # Restyle Word's real heading styles rather than faking headings with bold
    # Normal text. Semantics are the point: the navigation pane, screen readers, and
    # some parsers key off the style, not off the boldness. Word's stock headings are
    # blue Calibri Light, which would fight the design, so they get overridden here.
    for style_name, size in (("Heading 1", 22), ("Heading 2", 10)):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = normal.font.name
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb if style_name == "Heading 2" else RGBColor(0, 0, 0)
        pf = st.paragraph_format
        pf.space_before = Pt(10 if style_name == "Heading 2" else 0)
        pf.space_after = Pt(3 if style_name == "Heading 2" else 2)
        pf.keep_with_next = True          # a heading must never end a page

    b = data["basics"]
    h = doc.add_paragraph(b["name"], style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if b.get("headline"):
        p = doc.add_paragraph()
        rr = p.add_run(b["headline"])
        rr.font.size = Pt(11)
        rr.font.color.rgb = rgb
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)

    bits = [b.get("email"), b.get("phone"), b.get("location")]
    bits += [l.get("label") or l.get("url") for l in (b.get("links") or [])]
    p = doc.add_paragraph(" | ".join(str(x) for x in bits if x))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9.5)

    def heading(text: str):
        return doc.add_paragraph(text.upper(), style="Heading 2")

    if data.get("summary"):
        heading("Summary")
        doc.add_paragraph(str(data["summary"]).strip())

    if data.get("experience"):
        heading("Experience")
        for j in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{j.get('role','')}")
            r.bold = True
            if j.get("company"):
                p.add_run(f", {j['company']}")
            tail = " | ".join(x for x in (j.get("location"), date_range(j)) if x)
            if tail:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
                rr = sp.add_run(tail)
                rr.italic = True
                rr.font.size = Pt(9.5)
            for bl in (j.get("bullets") or []):
                doc.add_paragraph(str(bl), style="List Bullet")

    if data.get("projects"):
        heading("Projects")
        for e in data["projects"]:
            p = doc.add_paragraph()
            p.add_run(str(e.get("name", ""))).bold = True
            if e.get("summary"):
                p.add_run(f" — {e['summary']}")
            for bl in (e.get("bullets") or []):
                doc.add_paragraph(str(bl), style="List Bullet")

    if data.get("education"):
        heading("Education")
        for e in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(str(e.get("degree", ""))).bold = True
            tail = ", ".join(x for x in (e.get("institution"), date_range(e, "year")) if x)
            if tail:
                p.add_run(f" — {tail}")

    if data.get("skills"):
        heading("Skills")
        for g in data["skills"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            if isinstance(g, dict):
                p.add_run(f"{g.get('group','')}: ").bold = True
                p.add_run(", ".join(g.get("items") or []))
            else:
                p.add_run(str(g))

    for sec in (data.get("custom_sections") or []):
        heading(str(sec.get("title", "")))
        for it in (sec.get("items") or []):
            doc.add_paragraph(str(it if not isinstance(it, dict) else it.get("text", "")),
                              style="List Bullet")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


# --------------------------------------------------------------------- main

def main() -> int:
    ap = argparse.ArgumentParser(description="Build a resume from YAML.")
    ap.add_argument("data", help="path to resume.yaml")
    ap.add_argument("--template", "-t", default=None, help="template name (default: config.template or 'modern')")
    ap.add_argument("--format", "-f", default="pdf", help="comma list: pdf,html,docx,txt,all")
    ap.add_argument("--out", "-o", default=None, help="output directory (default: alongside the source)")
    ap.add_argument("--name", default=None, help="output basename (default: 'Firstname_Lastname_Resume')")
    ap.add_argument("--check", action="store_true", help="run ats_check.py on the produced PDF")
    ap.add_argument("--fit", action="store_true",
                    help="if the PDF exceeds config.max_pages, compress within tasteful "
                         "limits and report what was done")
    ap.add_argument("--quiet", "-q", action="store_true")
    a = ap.parse_args()

    src = Path(a.data).resolve()
    if not src.exists():
        sys.exit(f"error: no such file: {src}")
    data = load_resume(src)
    cfg = data.get("config", {})
    template = a.template or cfg.get("template", "modern")
    outdir = Path(a.out).resolve() if a.out else src.parent
    outdir.mkdir(parents=True, exist_ok=True)

    base = a.name or re.sub(r"[^\w]+", "_", data["basics"]["name"].strip()) + "_Resume"
    fmts = {f.strip().lower() for f in a.format.split(",") if f.strip()}
    if "all" in fmts:
        fmts = {"pdf", "html", "docx", "txt"}

    warnings = validate(data)
    if warnings and not a.quiet:
        print(f"content lint ({len(warnings)}):", file=sys.stderr)
        for w in warnings:
            print(f"  ! {w}", file=sys.stderr)
        print(file=sys.stderr)

    made: list[Path] = []
    html_path = outdir / f"{base}.html"
    max_pages = cfg.get("max_pages")

    if "pdf" in fmts:
        chrome = find_chrome()
        if not chrome:
            sys.exit("error: no Chrome/Chromium/Edge found. Set RESUME_CHROME=/path/to/chrome")
        pdf = outdir / f"{base}.pdf"
        if a.fit and max_pages:
            build_pdf_fitted(data, template, html_path, pdf, chrome, int(max_pages), a.quiet)
        else:
            html_path.write_text(render_html(data, template), encoding="utf-8")
            html_to_pdf(html_path, pdf, chrome)
            n = pdf_pages(pdf)
            if max_pages and n > int(max_pages) and not a.quiet:
                print(f"  ! {n} pages, but config.max_pages={max_pages}. "
                      f"Re-run with --fit to compress, or cut content.", file=sys.stderr)
        made.append(pdf)
        if "html" not in fmts:
            html_path.unlink(missing_ok=True)
    elif "html" in fmts:
        html_path.write_text(render_html(data, template), encoding="utf-8")

    if "html" in fmts and html_path.exists():
        made.append(html_path)

    if "docx" in fmts:
        d = outdir / f"{base}.docx"
        render_docx(data, d)
        made.append(d)

    if "txt" in fmts:
        t = outdir / f"{base}.txt"
        t.write_text(render_txt(data), encoding="utf-8")
        made.append(t)

    if not a.quiet:
        for m in made:
            size = m.stat().st_size
            print(f"  wrote {m}  ({size:,} bytes)")

    if a.check and (outdir / f"{base}.pdf").exists():
        print()
        rc = subprocess.run([sys.executable, str(HERE / "ats_check.py"),
                             str(outdir / f"{base}.pdf"), "--data", str(src)]).returncode
        return rc
    return 0


if __name__ == "__main__":
    sys.exit(main())
