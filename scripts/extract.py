#!/usr/bin/env python3
"""Extract an existing resume (PDF/DOCX/TXT) into editable resume.yaml.

    python extract.py old_resume.pdf -o resume.yaml

This is a STARTING POINT, not a conversion. Extraction from a PDF is lossy by
nature -- if the source had a sidebar or table layout, the text stream is already
interleaved before this script sees it (see ../DESIGN.md section 2), and no amount
of parsing recovers what the layout destroyed.

So: it does the boring 80% (contact details, section splitting, bullet capture),
marks what it is unsure about, and expects a human to fix the rest. It never
guesses a date it cannot read, and it never invents content.

Read the output back to the user and let them correct it.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
PHONE_RE = re.compile(r"(\+?\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}")
URL_RE = re.compile(r"(?:https?://)?(?:www\.)?((?:linkedin\.com|github\.com|gitlab\.com|"
                    r"behance\.net|dribbble\.com|medium\.com)/[\w./-]+)", re.I)
LOC_RE = re.compile(r"\b([A-Z][a-zA-Z.'-]+(?:\s[A-Z][a-zA-Z.'-]+)*),\s*"
                    r"([A-Z]{2}|[A-Z][a-z]+)\b")

# Month YYYY | MM/YYYY | YYYY, with an en/em/hyphen range and an open end.
MON = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*"
DATE = rf"(?:{MON}\.?\s+\d{{4}}|\d{{1,2}}/\d{{4}}|\d{{4}})"
RANGE_RE = re.compile(rf"({DATE})\s*[–—-]\s*({DATE}|Present|Current|Now)", re.I)

SECTION_ALIASES = {
    "experience": ["experience", "work experience", "professional experience",
                   "employment", "employment history", "work history", "career history",
                   "relevant experience", "professional background"],
    "education": ["education", "academic background", "academic qualifications",
                  "education & training", "qualifications"],
    "skills": ["skills", "technical skills", "core competencies", "competencies",
               "areas of expertise", "technologies", "expertise", "proficiencies"],
    "projects": ["projects", "selected projects", "personal projects", "side projects",
                 "portfolio", "key projects"],
    "certifications": ["certifications", "certificates", "licenses", "licences",
                       "certifications & licenses", "professional certifications"],
    "publications": ["publications", "papers", "selected publications", "research"],
    "awards": ["awards", "honors", "honours", "achievements", "awards & honors",
               "recognition"],
    "summary": ["summary", "professional summary", "profile", "about", "objective",
                "career summary", "executive summary", "overview"],
}
BULLET_CHARS = "•·▪◦‣⁃-–—*"


def read_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            sys.exit("error: pip install pdfplumber")
        with pdfplumber.open(path) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            sys.exit("error: pip install python-docx")
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:                       # sidebar content often hides here
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    if ext in (".txt", ".md", ".text"):
        return path.read_text(encoding="utf-8", errors="replace")
    sys.exit(f"error: unsupported file type '{ext}'. Use PDF, DOCX, or TXT.")


def classify(line: str) -> str | None:
    """Is this line a section heading?"""
    s = line.strip().strip(":").strip()
    if not s or len(s) > 42:
        return None
    low = re.sub(r"[^a-z& ]", "", s.lower()).strip()
    for key, names in SECTION_ALIASES.items():
        if low in names:
            return key
    # ALL CAPS short line with no sentence punctuation is very likely a heading
    if s.isupper() and len(s.split()) <= 4 and not s.endswith("."):
        for key, names in SECTION_ALIASES.items():
            if any(n in low for n in names):
                return key
    return None


def split_sections(text: str) -> tuple[list[str], dict[str, list[str]]]:
    head: list[str] = []
    out: dict[str, list[str]] = {}
    cur: str | None = None
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        key = classify(line)
        if key:
            cur = key
            out.setdefault(cur, [])
            continue
        (out[cur] if cur else head).append(line)
    return head, out


def parse_date(tok: str) -> str:
    tok = tok.strip().rstrip(".")
    if re.fullmatch(r"(?i)present|current|now", tok):
        return "present"
    m = re.fullmatch(rf"(?i)({MON})\.?\s+(\d{{4}})", tok)
    if m:
        mon = m.group(1)[:3].title()
        n = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
             "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"].index(mon) + 1
        return f"{m.group(2)}-{n:02d}"
    m = re.fullmatch(r"(\d{1,2})/(\d{4})", tok)
    if m:
        return f"{m.group(2)}-{int(m.group(1)):02d}"
    return tok


def parse_entries(lines: list[str]) -> list[dict]:
    """Group lines into entries. A dated line starts a new entry; bullets attach."""
    entries: list[dict] = []
    cur: dict | None = None
    for line in lines:
        s = line.strip()
        bullet = s[0] in BULLET_CHARS if s else False
        m = RANGE_RE.search(s)
        if m and not bullet:
            cur = {"_raw": s, "start": parse_date(m.group(1)), "end": parse_date(m.group(2)),
                   "bullets": []}
            head = RANGE_RE.sub("", s).strip(" ,·|–—-\t")
            loc = LOC_RE.search(head)
            if loc:
                cur["location"] = loc.group(0)
                head = head.replace(loc.group(0), "").strip(" ,·|–—-")
            parts = re.split(r"\s+[–—|·]\s+|,\s+(?=[A-Z])", head, maxsplit=1)
            cur["role"] = parts[0].strip() if parts else head
            if len(parts) > 1:
                cur["company"] = parts[1].strip()
            entries.append(cur)
        elif bullet and cur is not None:
            cur["bullets"].append(s.lstrip("".join(BULLET_CHARS)).strip())
        elif cur is not None and cur.get("bullets"):
            cur["bullets"][-1] += " " + s          # wrapped bullet
        elif cur is not None and not cur.get("company"):
            cur["company"] = s
        else:
            entries.append({"_raw": s, "_unparsed": True})
    return entries


def to_yaml(data: dict) -> str:
    import yaml

    class S(str): pass
    def rep(dumper, d):
        return dumper.represent_scalar("tag:yaml.org,2002:str", str(d),
                                       style="|" if "\n" in d else None)
    yaml.add_representer(S, rep)
    return yaml.dump(data, sort_keys=False, allow_unicode=True, width=88,
                     default_flow_style=False)


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract a resume into editable YAML.")
    ap.add_argument("source")
    ap.add_argument("-o", "--out", default="resume.yaml")
    ap.add_argument("--show-text", action="store_true", help="dump raw extracted text")
    a = ap.parse_args()

    src = Path(a.source).resolve()
    if not src.exists():
        sys.exit(f"error: no such file: {src}")
    text = read_text(src)
    if a.show_text:
        print(text)
        return 0
    if len(text.strip()) < 100:
        sys.exit("error: almost no text extracted — the source is probably a scan. "
                 "This script cannot OCR.")

    head, secs = split_sections(text)
    blob = "\n".join(head)

    basics: dict = {}
    lines = [l.strip() for l in head if l.strip()]
    if lines:
        basics["name"] = lines[0]
    if (m := EMAIL_RE.search(blob)):
        basics["email"] = m.group(0)
    if (m := PHONE_RE.search(blob)):
        basics["phone"] = m.group(0).strip()
    if (m := LOC_RE.search(blob)):
        basics["location"] = m.group(0)
    links = [{"label": u, "url": "https://" + u} for u in
             dict.fromkeys(m.group(1).rstrip("/.") for m in URL_RE.finditer(blob))]
    if links:
        basics["links"] = links

    out: dict = {
        "config": {"template": "modern", "page": "letter", "density": "normal",
                   "max_pages": 1},
        "basics": basics,
    }
    if secs.get("summary"):
        out["summary"] = " ".join(secs["summary"])
    for key in ("experience", "projects", "education"):
        if secs.get(key):
            ents = parse_entries(secs[key])
            cleaned = []
            for e in ents:
                e.pop("_raw", None)
                if e.pop("_unparsed", None) and not e.get("bullets"):
                    continue
                if key == "education":
                    e["degree"] = e.pop("role", "")
                    e["institution"] = e.pop("company", "")
                    e.pop("bullets", None) or None
                cleaned.append({k: v for k, v in e.items() if v})
            if cleaned:
                out[key] = cleaned
    if secs.get("skills"):
        groups = []
        for l in secs["skills"]:
            if ":" in l:
                g, items = l.split(":", 1)
                groups.append({"group": g.strip(),
                               "items": [x.strip() for x in re.split(r"[,;|·]", items) if x.strip()]})
            else:
                groups.append(l.strip())
        out["skills"] = groups
    for key in ("certifications", "awards", "publications"):
        if secs.get(key):
            out[key] = [l.strip().lstrip("".join(BULLET_CHARS)).strip() for l in secs[key]]

    banner = (
        "# Extracted by resume-designer/scripts/extract.py — REVIEW EVERY LINE.\n"
        "#\n"
        "# Extraction from a laid-out document is lossy. If the source used a sidebar,\n"
        "# columns, or tables, its text stream was already interleaved before this ran\n"
        "# (see DESIGN.md section 2) — some lines below may be spliced from two places.\n"
        "#\n"
        "# Nothing here was invented. Anything unreadable was dropped, not guessed.\n"
        "# Fix the content, then: python scripts/build.py resume.yaml --check\n\n"
    )
    Path(a.out).write_text(banner + to_yaml(out), encoding="utf-8")

    n_exp = len(out.get("experience") or [])
    print(f"  wrote {a.out}")
    print(f"  found: {len(basics)} contact field(s), {n_exp} experience entr{'y' if n_exp==1 else 'ies'}, "
          f"{len(out.get('education') or [])} education, {len(out.get('skills') or [])} skill group(s)")
    missing = [k for k in ("email", "phone", "location") if k not in basics]
    if missing:
        print(f"  ! could not find: {', '.join(missing)} — add by hand")
    if not out.get("experience"):
        print("  ! no dated experience entries recognised — the source layout was probably "
              "columnar. Read the raw text with --show-text and rebuild by hand.")
    print("  -> read it back to the user and correct it before building.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
